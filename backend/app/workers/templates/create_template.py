"""
Script to create a professional legal transcript DOCX template for docxtpl.

This creates a Word document template with proper formatting and Jinja2 placeholders
that will be populated by the TranscriptExporter class.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def create_transcript_template():
    """Create a professional legal transcript DOCX template."""
    doc = Document()

    # Set default font and margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    # Add custom styles
    styles = doc.styles

    # Header style
    if 'Transcript Header' not in styles:
        header_style = styles.add_style('Transcript Header', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.name = 'Arial'
        header_style.font.size = Pt(14)
        header_style.font.bold = True
        header_style.paragraph_format.space_after = Pt(12)

    # Subheader style
    if 'Transcript Subheader' not in styles:
        subheader_style = styles.add_style('Transcript Subheader', WD_STYLE_TYPE.PARAGRAPH)
        subheader_style.font.name = 'Arial'
        subheader_style.font.size = Pt(11)
        subheader_style.font.bold = True
        subheader_style.paragraph_format.space_after = Pt(6)

    # Body style
    if 'Transcript Body' not in styles:
        body_style = styles.add_style('Transcript Body', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing = 1.5

    # ========== DOCUMENT TITLE ==========
    title = doc.add_paragraph('LEGAL TRANSCRIPT', style='Transcript Header')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Blank line

    # ========== CASE INFORMATION ==========
    case_header = doc.add_paragraph('CASE INFORMATION', style='Transcript Subheader')
    case_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Case details table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    # Case Number
    table.cell(0, 0).text = 'Case Number:'
    table.cell(0, 1).text = '{{ case_number }}'

    # Case Name
    table.cell(1, 0).text = 'Case Name:'
    table.cell(1, 1).text = '{{ case_name }}'

    # Client
    table.cell(2, 0).text = 'Client:'
    table.cell(2, 1).text = '{{ client }}'

    # Matter Type
    table.cell(3, 0).text = 'Matter Type:'
    table.cell(3, 1).text = '{{ matter_type }}'

    # Proceeding Date
    table.cell(4, 0).text = 'Proceeding Date:'
    table.cell(4, 1).text = '{{ proceeding_date }}'

    # Location
    table.cell(5, 0).text = 'Location:'
    table.cell(5, 1).text = '{{ location }}'

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)

    doc.add_paragraph()  # Blank line

    # ========== DOCUMENT INFORMATION ==========
    doc_header = doc.add_paragraph('DOCUMENT INFORMATION', style='Transcript Subheader')

    doc_table = doc.add_table(rows=4, cols=2)
    doc_table.style = 'Light Grid Accent 1'

    doc_table.cell(0, 0).text = 'Document Name:'
    doc_table.cell(0, 1).text = '{{ document_name }}'

    doc_table.cell(1, 0).text = 'Witness/Deponent:'
    doc_table.cell(1, 1).text = '{{ witness }}'

    doc_table.cell(2, 0).text = 'Total Duration:'
    doc_table.cell(2, 1).text = '{{ total_duration }}'

    doc_table.cell(3, 0).text = 'Segment Count:'
    doc_table.cell(3, 1).text = '{{ segment_count }}'

    for row in doc_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)

    doc.add_paragraph()  # Blank line

    # ========== PARTICIPANTS/SPEAKERS ==========
    speakers_header = doc.add_paragraph('PARTICIPANTS', style='Transcript Subheader')

    # Speakers table with Jinja2 loop
    doc.add_paragraph('{% if speakers %}')

    speakers_text = doc.add_paragraph()
    speakers_text.add_run('{% for speaker in speakers %}')

    speaker_item = doc.add_paragraph(style='List Bullet')
    speaker_item.add_run('{{ speaker.label }}{% if speaker.role %} - {{ speaker.role }}{% endif %}{% if speaker.affiliation %} ({{ speaker.affiliation }}){% endif %}')

    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')

    doc.add_paragraph()  # Blank line

    # ========== TRANSCRIPT HEADER ==========
    transcript_header = doc.add_paragraph('TRANSCRIPT', style='Transcript Header')
    transcript_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Blank line

    # ========== TRANSCRIPT SEGMENTS ==========
    # Add instruction comment (will be removed in final render)
    doc.add_paragraph('{% for segment in segments %}')

    # Speaker label and timestamp
    segment_header = doc.add_paragraph()
    segment_header.add_run('{{ segment.speaker_label }}').bold = True
    segment_header.add_run(' [{{ segment.timestamp }}]')
    segment_header.paragraph_format.space_before = Pt(8)
    segment_header.paragraph_format.space_after = Pt(4)

    # Segment text
    segment_text = doc.add_paragraph('{{ segment.text }}', style='Transcript Body')
    segment_text.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph('{% endfor %}')

    doc.add_paragraph()  # Blank line
    doc.add_page_break()

    # ========== FOOTER/CERTIFICATION ==========
    footer_header = doc.add_paragraph('CERTIFICATION', style='Transcript Subheader')

    certification = doc.add_paragraph(
        'This transcript was generated electronically from audio/video recording on {{ export_date }} '
        'at {{ export_time }}. The transcript represents the best effort to accurately capture the '
        'spoken content from the recording.'
    )
    certification.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()  # Blank line

    signature_line = doc.add_paragraph()
    signature_line.add_run('_' * 50)
    signature_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

    signature_label = doc.add_paragraph('Certified By')
    signature_label.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    date_line = doc.add_paragraph()
    date_line.add_run('Date: _' * 20)

    # Save the template
    template_path = Path(__file__).parent / 'transcript_template.docx'
    doc.save(str(template_path))
    print(f"Template created successfully at: {template_path}")
    return template_path


if __name__ == '__main__':
    create_transcript_template()
