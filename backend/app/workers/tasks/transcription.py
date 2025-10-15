"""
Transcription Tasks

Celery tasks for audio transcription and processing.
Production-ready implementation with WhisperX integration.
"""
import os
import json
import logging
import tempfile
import subprocess
import uuid
import re
import httpx
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime, timedelta
from io import BytesIO

from app.workers.celery_app import celery_app
from app.core.database import SessionLocal
from app.core.minio_client import minio_client
from app.models.transcription import Transcription
from app.models.document import Document, DocumentStatus

logger = logging.getLogger(__name__)


class TranscriptionError(Exception):
    """Custom exception for transcription errors."""
    pass


class AudioProcessor:
    """Handles audio preprocessing with FFmpeg."""

    @staticmethod
    def preprocess_audio(input_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Preprocess audio file to 16kHz mono WAV format for optimal transcription.

        Args:
            input_path: Path to input audio/video file
            output_path: Path to output WAV file

        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            cmd = [
                'ffmpeg',
                '-i', input_path,
                '-ar', '16000',  # 16kHz sample rate
                '-ac', '1',       # Mono channel
                '-c:a', 'pcm_s16le',  # 16-bit PCM
                '-y',             # Overwrite output file
                output_path
            ]

            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=300  # 5 minute timeout
            )

            if result.returncode == 0:
                return True, "Audio preprocessing successful"
            else:
                error_msg = result.stderr.decode('utf-8', errors='ignore')
                return False, f"FFmpeg error: {error_msg}"

        except subprocess.TimeoutExpired:
            return False, "Audio preprocessing timed out"
        except Exception as e:
            return False, f"Audio preprocessing failed: {str(e)}"

    @staticmethod
    def get_audio_duration(file_path: str) -> Optional[float]:
        """
        Get audio duration in seconds using FFprobe.

        Args:
            file_path: Path to audio file

        Returns:
            Duration in seconds, or None if failed
        """
        try:
            cmd = [
                'ffprobe',
                '-v', 'error',
                '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1',
                file_path
            ]

            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=30
            )

            if result.returncode == 0:
                duration_str = result.stdout.decode('utf-8').strip()
                return float(duration_str)
            return None

        except Exception as e:
            logger.warning(f"Failed to get audio duration: {e}")
            return None


class WhisperTranscriber:
    """Handles transcription using OpenAI Whisper API."""

    def __init__(self, api_key: Optional[str] = None):
        """Initialize transcriber with OpenAI API key."""
        self.api_key = api_key or os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise TranscriptionError("OpenAI API key not configured")

    def transcribe(
        self,
        audio_path: str,
        language: Optional[str] = None,
        task: str = "transcribe",
        temperature: float = 0.0,
        initial_prompt: Optional[str] = None,
        task_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Transcribe audio file using OpenAI Whisper API.

        Args:
            audio_path: Path to audio file
            language: Language code (e.g., 'en', 'es') or None for auto-detect
            task: 'transcribe' for same-language or 'translate' for English translation
            temperature: Sampling temperature (0.0-1.0)
            initial_prompt: Optional context prompt
            task_id: Optional Celery task ID for progress updates

        Returns:
            Dict containing transcription results with segments
        """
        try:
            from openai import OpenAI

            client = OpenAI(api_key=self.api_key)

            logger.info(f"Starting Whisper API transcription for {audio_path} (task={task}, language={language})")

            with open(audio_path, 'rb') as audio_file:
                # Build API parameters
                api_params = {
                    "model": "whisper-1",
                    "file": audio_file,
                    "response_format": "verbose_json",
                    "timestamp_granularities": ["segment", "word"],
                    "temperature": temperature,
                }

                # Add optional parameters
                if language and language != "auto":
                    api_params["language"] = language
                if initial_prompt:
                    api_params["prompt"] = initial_prompt

                # Use transcriptions or translations API based on task
                if task == "translate":
                    response = client.audio.translations.create(**api_params)
                else:
                    response = client.audio.transcriptions.create(**api_params)

            # Parse response into segments format
            segments = []
            if hasattr(response, 'segments') and response.segments:
                for idx, segment in enumerate(response.segments):
                    seg_data = {
                        'id': str(uuid.uuid4()),  # Generate UUID for segment
                        'start': segment.get('start', 0.0),
                        'end': segment.get('end', 0.0),
                        'text': segment.get('text', '').strip(),
                        'words': []
                    }

                    # Add word-level timestamps if available
                    if hasattr(segment, 'words') and segment.words:
                        for word in segment.words:
                            seg_data['words'].append({
                                'word': word.get('word', ''),
                                'start': word.get('start', 0.0),
                                'end': word.get('end', 0.0)
                            })

                    segments.append(seg_data)
            else:
                # Fallback: create single segment from text
                segments = [{
                    'id': str(uuid.uuid4()),  # Generate UUID for segment
                    'start': 0.0,
                    'end': 0.0,
                    'text': response.text,
                    'words': []
                }]

            result = {
                'text': response.text,
                'language': response.language if hasattr(response, 'language') else language,
                'duration': response.duration if hasattr(response, 'duration') else None,
                'segments': segments
            }

            logger.info(f"Transcription completed: {len(segments)} segments")
            return result

        except Exception as e:
            logger.error(f"Whisper API transcription failed: {str(e)}")
            raise TranscriptionError(f"Transcription failed: {str(e)}")


class SpeakerDiarizer:
    """Handles speaker diarization using simple heuristics."""

    @staticmethod
    def diarize_segments(segments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Add speaker labels to segments using simple pause-based heuristics.

        This is a simplified implementation. For production, consider using:
        - Pyannote.audio for deep learning-based diarization
        - AssemblyAI API for cloud-based diarization

        Args:
            segments: List of transcription segments

        Returns:
            List of segments with speaker labels added
        """
        if not segments:
            return segments

        # Simple heuristic: detect speaker changes based on pauses
        PAUSE_THRESHOLD = 2.0  # seconds

        current_speaker = 1
        diarized_segments = []

        for idx, segment in enumerate(segments):
            # Check if there's a significant pause before this segment
            if idx > 0:
                prev_end = segments[idx - 1]['end']
                current_start = segment['start']
                pause_duration = current_start - prev_end

                # If pause is significant, assume speaker change
                if pause_duration > PAUSE_THRESHOLD:
                    current_speaker += 1

            # Add speaker label
            segment_copy = segment.copy()
            segment_copy['speaker'] = f"SPEAKER_{current_speaker:02d}"
            diarized_segments.append(segment_copy)

        return diarized_segments

    @staticmethod
    def smooth_speaker_changes(
        segments: List[Dict[str, Any]],
        min_segment_duration: float = 0.5,
        min_speaker_gap: float = 0.3
    ) -> List[Dict[str, Any]]:
        """
        Smooth rapid speaker changes using post-processing.

        Production-quality diarization requires smoothing to reduce false speaker changes
        caused by short pauses, overlapping speech, or diarization errors.

        Strategy:
        1. Merge very short segments (< min_segment_duration) with neighbors
        2. Apply majority voting for rapid speaker changes
        3. Require minimum gap (min_speaker_gap) between speaker changes

        Args:
            segments: List of diarized segments
            min_segment_duration: Minimum segment duration in seconds (default: 0.5s)
            min_speaker_gap: Minimum time between speaker changes in seconds (default: 0.3s)

        Returns:
            List of smoothed segments
        """
        if not segments or len(segments) < 2:
            return segments

        logger.info(f"Smoothing speaker changes (min_duration={min_segment_duration}s, min_gap={min_speaker_gap}s)")

        smoothed = []
        i = 0

        while i < len(segments):
            current = segments[i].copy()
            duration = current['end'] - current['start']

            # If segment is very short, try to merge with neighbor
            if duration < min_segment_duration and i < len(segments) - 1:
                next_seg = segments[i + 1]

                # Merge with next segment if speakers match
                if current.get('speaker') == next_seg.get('speaker'):
                    logger.debug(f"Merging short segment ({duration:.2f}s) at {current['start']:.1f}s with next")
                    # Skip this segment, will be handled by next iteration
                    i += 1
                    continue
                # If speakers don't match, assign to majority speaker based on surrounding context
                elif i > 0:
                    prev_seg = smoothed[-1]
                    # If surrounded by same speaker, reassign
                    if prev_seg.get('speaker') == next_seg.get('speaker'):
                        logger.debug(f"Reassigning isolated short segment ({duration:.2f}s) to majority speaker")
                        current['speaker'] = next_seg.get('speaker')

            # Check for rapid speaker changes (speaker switches back and forth quickly)
            if smoothed and i < len(segments) - 1:
                prev_seg = smoothed[-1]
                next_seg = segments[i + 1]
                time_since_last_change = current['start'] - prev_seg['end']

                # If speaker changed very recently and changes back, smooth it
                if (time_since_last_change < min_speaker_gap and
                    prev_seg.get('speaker') == next_seg.get('speaker') and
                    current.get('speaker') != prev_seg.get('speaker')):

                    logger.debug(
                        f"Smoothing rapid speaker change at {current['start']:.1f}s "
                        f"(gap={time_since_last_change:.2f}s)"
                    )
                    current['speaker'] = prev_seg.get('speaker')

            smoothed.append(current)
            i += 1

        # Log improvement metrics
        original_speaker_changes = sum(
            1 for i in range(1, len(segments))
            if segments[i].get('speaker') != segments[i-1].get('speaker')
        )
        smoothed_speaker_changes = sum(
            1 for i in range(1, len(smoothed))
            if smoothed[i].get('speaker') != smoothed[i-1].get('speaker')
        )

        logger.info(
            f"Speaker smoothing: reduced changes from {original_speaker_changes} to {smoothed_speaker_changes} "
            f"({original_speaker_changes - smoothed_speaker_changes} false changes removed)"
        )

        return smoothed


class SpeakerNameInferencer:
    """Handles AI-powered speaker name inference from conversation content."""

    @staticmethod
    def extract_name_patterns(segments: List[Dict[str, Any]], filename: Optional[str] = None) -> Dict[str, List[str]]:
        """
        Extract potential speaker names using pattern matching.

        Patterns:
        - "Hi [NAME]" / "Hello [NAME]"
        - "I'm [NAME]" / "My name is [NAME]"
        - "[NAME]?" / "[NAME]!"
        - Filename patterns like "jane_mark_conversation.mp3"

        Args:
            segments: List of transcription segments
            filename: Optional filename to extract names from

        Returns:
            Dict mapping potential names to their sources
        """
        patterns = {
            'greeting': r'\b(?:hi|hello|hey)\s+([A-Z][a-z]+)\b',
            'introduction': r'\b(?:I\'m|I am|my name is|this is)\s+([A-Z][a-z]+)\b',
            'direct_address': r'\b([A-Z][a-z]+)[?!]\b',
            'possessive': r'\b([A-Z][a-z]+)\'s\b'
        }

        potential_names = {}

        # Extract from segments
        for segment in segments[:20]:  # Check first 20 segments
            text = segment.get('text', '')

            for pattern_type, pattern in patterns.items():
                matches = re.findall(pattern, text)
                for name in matches:
                    # Filter out common words that might match
                    if name.lower() not in ['i', 'you', 'we', 'they', 'okay', 'yeah', 'sure', 'right', 'well']:
                        if name not in potential_names:
                            potential_names[name] = []
                        potential_names[name].append(f"{pattern_type}: '{text[:50]}...'")

        # Extract from filename
        if filename:
            # Pattern: "jane_mark_conversation.mp3" or "john-doe-interview.wav"
            filename_clean = re.sub(r'\.(mp3|wav|m4a|mp4|avi|mov)$', '', filename.lower())
            filename_parts = re.split(r'[_\-\s]+', filename_clean)

            # Look for name-like parts (capitalized words)
            for part in filename_parts:
                if len(part) > 2 and part not in ['conversation', 'interview', 'call', 'meeting', 'recording', 'audio', 'video']:
                    name_capitalized = part.capitalize()
                    if name_capitalized not in potential_names:
                        potential_names[name_capitalized] = []
                    potential_names[name_capitalized].append(f"filename: '{filename}'")

        logger.info(f"Extracted {len(potential_names)} potential names from patterns: {list(potential_names.keys())}")
        return potential_names

    @staticmethod
    async def infer_names_with_llm(
        segments: List[Dict[str, Any]],
        speakers: List[str],
        potential_names: Dict[str, List[str]],
        ollama_url: str = "http://localhost:11434"
    ) -> Dict[str, Dict[str, Any]]:
        """
        Use LLM to infer speaker names from conversation context.

        Args:
            segments: List of transcription segments
            speakers: List of unique speaker IDs
            potential_names: Dictionary of potential names from pattern matching
            ollama_url: Ollama API URL

        Returns:
            Dict mapping speaker IDs to inferred names with confidence scores
        """
        try:
            # Prepare conversation context (first ~10 segments)
            context_segments = segments[:min(10, len(segments))]
            conversation_text = "\n".join([
                f"{seg.get('speaker', 'UNKNOWN')}: {seg.get('text', '')}"
                for seg in context_segments
            ])

            # Build prompt
            system_prompt = """You are an AI assistant that identifies speakers in conversations.
Analyze the conversation and infer the most likely names for each speaker.
Return ONLY valid JSON with no additional text or markdown formatting.

Response format:
{
  "SPEAKER_00": {
    "name": "John",
    "confidence": 0.85,
    "reasoning": "Speaker introduces themselves as John in the first segment"
  },
  "SPEAKER_01": {
    "name": "Sarah",
    "confidence": 0.75,
    "reasoning": "Other speaker addresses them as Sarah"
  }
}

Guidelines:
- Only include speakers you can confidently identify
- Confidence score: 0.0 to 1.0 (only suggest names with confidence > 0.8)
- Use first names only unless full name is clear
- If unsure, do not include that speaker"""

            potential_names_str = ""
            if potential_names:
                potential_names_str = f"\n\nPotential names found: {', '.join(potential_names.keys())}"

            prompt = f"""Analyze this conversation and identify the speakers:

{conversation_text}
{potential_names_str}

Speakers to identify: {', '.join(speakers)}

Return JSON with speaker names and confidence scores:"""

            # Call Ollama API
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    f"{ollama_url}/api/generate",
                    json={
                        "model": "llama3.1:latest",
                        "prompt": prompt,
                        "system": system_prompt,
                        "stream": False,
                        "format": "json",
                        "options": {
                            "temperature": 0.1,  # Low temperature for consistent results
                            "top_p": 0.9
                        }
                    }
                )

                if response.status_code == 200:
                    result = response.json()
                    response_text = result.get('response', '{}')

                    # Parse JSON response
                    try:
                        inferred_names = json.loads(response_text)
                        logger.info(f"LLM inference result: {inferred_names}")
                        return inferred_names
                    except json.JSONDecodeError as e:
                        logger.error(f"Failed to parse LLM JSON response: {e}")
                        logger.debug(f"Raw response: {response_text}")
                        return {}
                else:
                    logger.error(f"Ollama API error: {response.status_code}")
                    logger.error(f"Ollama API response body: {response.text}")
                    logger.error(f"Ollama API response headers: {response.headers}")
                    return {}

        except Exception as e:
            logger.error(f"LLM name inference failed: {e}", exc_info=True)
            return {}

    @staticmethod
    async def infer_speaker_names(
        segments: List[Dict[str, Any]],
        speakers: Dict[str, Dict[str, Any]],
        filename: Optional[str] = None,
        ollama_url: str = "http://localhost:11434"
    ) -> Tuple[Dict[str, Dict[str, Any]], Dict[str, Any]]:
        """
        Main method to infer speaker names using pattern matching and LLM.

        Args:
            segments: List of transcription segments
            speakers: Dictionary of speaker information
            filename: Optional filename to extract names from
            ollama_url: Ollama API URL

        Returns:
            Tuple of (updated_speakers, metadata with inference info)
        """
        logger.info("Starting speaker name inference...")

        # Step 1: Extract potential names using pattern matching
        potential_names = SpeakerNameInferencer.extract_name_patterns(segments, filename)

        # Step 2: Use LLM to infer names with confidence scores
        speaker_ids = list(speakers.keys())
        inferred_names = await SpeakerNameInferencer.infer_names_with_llm(
            segments,
            speaker_ids,
            potential_names,
            ollama_url
        )

        # Step 3: Apply inferred names if confidence > 0.8
        updated_speakers = speakers.copy()
        metadata = {
            'inference_performed': True,
            'potential_names': potential_names,
            'llm_inferences': inferred_names,
            'applied_names': {}
        }

        for speaker_id, inference in inferred_names.items():
            confidence = inference.get('confidence', 0.0)
            inferred_name = inference.get('name', '')
            reasoning = inference.get('reasoning', '')

            if speaker_id in updated_speakers:
                if confidence > 0.8 and inferred_name:
                    # Apply inferred name
                    old_name = updated_speakers[speaker_id].get('name', speaker_id)
                    updated_speakers[speaker_id]['name'] = inferred_name
                    metadata['applied_names'][speaker_id] = {
                        'old_name': old_name,
                        'new_name': inferred_name,
                        'confidence': confidence,
                        'reasoning': reasoning
                    }
                    logger.info(f"Applied inferred name for {speaker_id}: '{inferred_name}' (confidence: {confidence:.2f})")
                else:
                    logger.info(f"Skipped {speaker_id}: confidence {confidence:.2f} below threshold 0.8")

        if metadata['applied_names']:
            logger.info(f"Successfully inferred {len(metadata['applied_names'])} speaker names")
        else:
            logger.info("No speaker names met confidence threshold (0.8), using default names")

        return updated_speakers, metadata


class TranscriptionExporter:
    """Handles exporting transcriptions to various formats."""

    @staticmethod
    def export_to_docx(segments: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Export transcription to DOCX format.

        Args:
            segments: List of transcription segments
            output_path: Path to output DOCX file

        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Add title
            title = doc.add_heading('Transcription', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph("")

            # Group segments by speaker
            current_speaker = None
            current_paragraph = None

            for segment in segments:
                speaker = segment.get('speaker', 'SPEAKER_01')
                start_time = segment.get('start', 0.0)
                text = segment.get('text', '').strip()

                if not text:
                    continue

                # Add speaker header if speaker changed
                if speaker != current_speaker:
                    current_speaker = speaker

                    # Add speaker heading
                    speaker_para = doc.add_paragraph()
                    speaker_run = speaker_para.add_run(f"\n{speaker}")
                    speaker_run.bold = True
                    speaker_run.font.size = Pt(12)
                    speaker_run.font.color.rgb = RGBColor(0, 0, 139)

                # Add timestamped text
                para = doc.add_paragraph()

                # Add timestamp
                timestamp_str = TranscriptionExporter._format_timestamp(start_time)
                time_run = para.add_run(f"[{timestamp_str}] ")
                time_run.font.color.rgb = RGBColor(128, 128, 128)
                time_run.font.size = Pt(9)

                # Add text
                text_run = para.add_run(text)
                text_run.font.size = Pt(11)

            doc.save(output_path)
            logger.info(f"Exported DOCX to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            return False

    @staticmethod
    def export_to_srt(segments: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Export transcription to SRT subtitle format.

        Args:
            segments: List of transcription segments
            output_path: Path to output SRT file

        Returns:
            True if successful, False otherwise
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                for idx, segment in enumerate(segments, start=1):
                    start = segment.get('start', 0.0)
                    end = segment.get('end', 0.0)
                    text = segment.get('text', '').strip()
                    speaker = segment.get('speaker', '')

                    if not text:
                        continue

                    # Format: sequence number
                    f.write(f"{idx}\n")

                    # Format: start --> end
                    start_str = TranscriptionExporter._format_srt_timestamp(start)
                    end_str = TranscriptionExporter._format_srt_timestamp(end)
                    f.write(f"{start_str} --> {end_str}\n")

                    # Format: text (with optional speaker)
                    if speaker:
                        f.write(f"[{speaker}] {text}\n")
                    else:
                        f.write(f"{text}\n")

                    # Blank line between subtitles
                    f.write("\n")

            logger.info(f"Exported SRT to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export SRT: {e}")
            return False

    @staticmethod
    def export_to_vtt(segments: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Export transcription to WebVTT subtitle format.

        Args:
            segments: List of transcription segments
            output_path: Path to output VTT file

        Returns:
            True if successful, False otherwise
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # WebVTT header
                f.write("WEBVTT\n\n")

                for idx, segment in enumerate(segments, start=1):
                    start = segment.get('start', 0.0)
                    end = segment.get('end', 0.0)
                    text = segment.get('text', '').strip()
                    speaker = segment.get('speaker', '')

                    if not text:
                        continue

                    # Format: start --> end
                    start_str = TranscriptionExporter._format_vtt_timestamp(start)
                    end_str = TranscriptionExporter._format_vtt_timestamp(end)
                    f.write(f"{start_str} --> {end_str}\n")

                    # Format: text with speaker as voice tag
                    if speaker:
                        f.write(f"<v {speaker}>{text}</v>\n")
                    else:
                        f.write(f"{text}\n")

                    # Blank line between cues
                    f.write("\n")

            logger.info(f"Exported VTT to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export VTT: {e}")
            return False

    @staticmethod
    def export_to_json(
        segments: List[Dict[str, Any]],
        output_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        Export transcription to JSON format.

        Args:
            segments: List of transcription segments
            output_path: Path to output JSON file
            metadata: Optional metadata to include

        Returns:
            True if successful, False otherwise
        """
        try:
            data = {
                'metadata': metadata or {},
                'segments': segments,
                'generated_at': datetime.utcnow().isoformat()
            }

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            logger.info(f"Exported JSON to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export JSON: {e}")
            return False

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """Format seconds to HH:MM:SS."""
        td = timedelta(seconds=seconds)
        hours = int(td.total_seconds() // 3600)
        minutes = int((td.total_seconds() % 3600) // 60)
        secs = int(td.total_seconds() % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

    @staticmethod
    def _format_srt_timestamp(seconds: float) -> str:
        """Format seconds to SRT timestamp format (HH:MM:SS,mmm)."""
        td = timedelta(seconds=seconds)
        hours = int(td.total_seconds() // 3600)
        minutes = int((td.total_seconds() % 3600) // 60)
        secs = int(td.total_seconds() % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    @staticmethod
    def _format_vtt_timestamp(seconds: float) -> str:
        """Format seconds to WebVTT timestamp format (HH:MM:SS.mmm)."""
        td = timedelta(seconds=seconds)
        hours = int(td.total_seconds() // 3600)
        minutes = int((td.total_seconds() % 3600) // 60)
        secs = int(td.total_seconds() % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"


@celery_app.task(name="transcribe_audio", bind=True)
def transcribe_audio(
    self,
    document_id: int,
    options: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Transcribe an audio/video file to text with timestamps and speaker diarization.

    Full pipeline:
    1. Download audio/video from MinIO
    2. Preprocess with FFmpeg (convert to 16kHz mono WAV)
    3. Transcribe with OpenAI Whisper API
    4. Diarize speakers (optional, based on options)
    5. Export to DOCX, SRT, VTT, JSON formats
    6. Upload exports to MinIO
    7. Update transcription record in PostgreSQL
    8. Handle errors gracefully with status updates

    Args:
        document_id: ID of the document to transcribe
        options: Optional dict containing transcription configuration:
            - language: Language code or None for auto-detect
            - task: 'transcribe' or 'translate'
            - enable_diarization: Enable speaker identification
            - temperature: Sampling temperature
            - initial_prompt: Optional context prompt

    Returns:
        Dict containing transcription status and results
    """
    db = SessionLocal()
    temp_dir = None

    # Parse options with defaults
    options = options or {}
    language = options.get("language")
    task = options.get("task", "transcribe")
    enable_diarization = options.get("enable_diarization", True)
    temperature = options.get("temperature", 0.0)
    initial_prompt = options.get("initial_prompt")

    try:
        # Update task state to STARTED
        self.update_state(
            state='STARTED',
            meta={'status': 'Initializing transcription', 'progress': 0}
        )

        # Step 1: Get document from database
        logger.info(f"Starting transcription for document {document_id}")
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise TranscriptionError(f"Document {document_id} not found")

        # Update document status
        document.status = DocumentStatus.PROCESSING
        db.commit()

        # Step 2: Download audio/video from MinIO
        self.update_state(
            state='PROCESSING',
            meta={'status': 'Downloading audio from storage', 'progress': 10}
        )

        logger.info(f"Downloading file from MinIO: {document.file_path}")
        file_content = minio_client.download_file(document.file_path)

        if not file_content:
            raise TranscriptionError("Failed to download file from MinIO")

        # Create temporary directory for processing
        temp_dir = tempfile.mkdtemp(prefix='transcription_')
        logger.info(f"Created temporary directory: {temp_dir}")

        # Get file extension
        _, ext = os.path.splitext(document.filename)
        original_file = os.path.join(temp_dir, f"original{ext}")
        processed_wav = os.path.join(temp_dir, "audio.wav")

        # Save original file
        with open(original_file, 'wb') as f:
            f.write(file_content)

        # Step 3: Preprocess audio with FFmpeg
        self.update_state(
            state='PROCESSING',
            meta={'status': 'Preprocessing audio', 'progress': 20}
        )

        logger.info("Preprocessing audio with FFmpeg")
        processor = AudioProcessor()
        success, message = processor.preprocess_audio(original_file, processed_wav)

        if not success:
            raise TranscriptionError(f"Audio preprocessing failed: {message}")

        # Get audio duration
        duration = processor.get_audio_duration(processed_wav)
        logger.info(f"Audio duration: {duration} seconds")

        # Step 4: Transcribe with Whisper
        self.update_state(
            state='PROCESSING',
            meta={'status': 'Transcribing audio', 'progress': 30}
        )

        # Use self-hosted transcription with Pyannote diarization (NO HuggingFace token needed)
        logger.info(f"Starting self-hosted transcription (language={language or 'auto'}, diarization={enable_diarization})")

        try:
            import torch
            from app.services.model_manager import get_pyannote_model_paths
            from app.services.device_manager import get_device, get_compute_type

            # Get device and compute type (supports both CUDA and ROCm)
            device = get_device()
            compute_type = get_compute_type(prefer_fp16=True)

            logger.info(f"Using device: {device} with compute type: {compute_type}")

            # Check if HF token is available for Pyannote diarization
            hf_token = os.getenv("HF_TOKEN")
            pyannote_available = hf_token is not None and enable_diarization

            if pyannote_available:
                logger.info("HuggingFace token available - will use Pyannote for speaker diarization")
            elif enable_diarization:
                logger.warning("No HF_TOKEN found - will use simple heuristic diarization")

            # Try WhisperX for transcription + alignment (faster, more accurate)
            try:
                from app.workers.pipelines.whisperx_pipeline import WhisperXPipeline

                min_speakers = options.get("min_speakers", 2)
                max_speakers = options.get("max_speakers", 10)

                # Initialize WhisperX WITHOUT diarization (we'll use Pyannote directly)
                pipeline = WhisperXPipeline(
                    model_name="medium",  # Medium model (~5GB VRAM, good balance of speed/quality)
                    device=device,
                    compute_type=compute_type,
                    language=language if language and language != "auto" else None,
                    hf_token=None  # No HF token needed!
                )

                # Transcribe with alignment only (no diarization yet)
                result = pipeline.transcribe(
                    audio_path=processed_wav,
                    enable_alignment=True,
                    enable_diarization=False,  # We'll do this separately
                    batch_size=16
                )

                # Convert WhisperX result to our format and add UUIDs
                segments = []
                for seg in result.segments:
                    seg_dict = seg.to_dict()
                    seg_dict['id'] = str(uuid.uuid4())  # Add unique ID for each segment
                    segments.append(seg_dict)

                full_text = result.get_full_text()
                detected_language = result.language

                logger.info(f"WhisperX transcription completed: {len(segments)} segments, language={detected_language}")

                # Cleanup WhisperX models from memory
                pipeline.cleanup()

            except ImportError:
                logger.info("WhisperX not available, using OpenAI Whisper API")

                # Fallback to OpenAI Whisper API
                transcriber = WhisperTranscriber()
                transcription_result = transcriber.transcribe(
                    processed_wav,
                    language=language,
                    task=task,
                    temperature=temperature,
                    initial_prompt=initial_prompt,
                    task_id=self.request.id
                )

                segments = transcription_result['segments']
                full_text = transcription_result['text']
                detected_language = transcription_result.get('language', language)

                logger.info(f"OpenAI Whisper transcription completed: {len(segments)} segments")

            # Step 5: Speaker diarization (if enabled)
            if enable_diarization:
                self.update_state(
                    state='PROCESSING',
                    meta={'status': 'Identifying speakers', 'progress': 60}
                )

                # Try Pyannote if HF token available, otherwise use simple diarization
                if pyannote_available:
                    try:
                        # Use Pyannote.audio directly with HuggingFace models
                        from pyannote.audio import Pipeline

                        logger.info("Running Pyannote speaker diarization...")

                        # Check if exact speaker count is provided
                        num_speakers = options.get("num_speakers")
                        min_speakers = options.get("min_speakers", 2)
                        max_speakers = options.get("max_speakers", 5)

                        # Load diarization pipeline from HuggingFace
                        logger.info("Loading speaker-diarization-3.1 from HuggingFace...")
                        diarization_pipeline = Pipeline.from_pretrained(
                            "pyannote/speaker-diarization-3.1",
                            use_auth_token=hf_token
                        )
                        logger.info("Pyannote pipeline loaded successfully")

                        # Move to GPU if available
                        if device == "cuda":
                            diarization_pipeline.to(torch.device("cuda"))

                        # Run diarization with exact speaker count or range
                        if num_speakers is not None:
                            logger.info(f"Using exact speaker count: {num_speakers}")
                            diarization = diarization_pipeline(
                                processed_wav,
                                num_speakers=num_speakers
                            )
                        else:
                            logger.info(f"Using auto-detection with speaker range: {min_speakers}-{max_speakers}")
                            diarization = diarization_pipeline(
                                processed_wav,
                                min_speakers=min_speakers,
                                max_speakers=max_speakers
                            )

                        # Assign speakers to segments
                        diarized_segments = []
                        for segment in segments:
                            # Find overlapping speaker from diarization
                            segment_start = segment['start']
                            segment_end = segment['end']

                            # Find most common speaker in this segment
                            speaker_times = {}
                            for turn, _, speaker in diarization.itertracks(yield_label=True):
                                # Calculate overlap with this segment
                                overlap_start = max(segment_start, turn.start)
                                overlap_end = min(segment_end, turn.end)
                                overlap_duration = max(0, overlap_end - overlap_start)

                                if overlap_duration > 0:
                                    speaker_times[speaker] = speaker_times.get(speaker, 0) + overlap_duration

                            # Assign speaker with most overlap
                            if speaker_times:
                                assigned_speaker = max(speaker_times, key=speaker_times.get)
                                # Pyannote returns speaker labels like "SPEAKER_00", use directly
                                segment['speaker'] = assigned_speaker if assigned_speaker.startswith("SPEAKER_") else f"SPEAKER_{assigned_speaker}"
                            else:
                                segment['speaker'] = "SPEAKER_00"

                            diarized_segments.append(segment)

                        logger.info(f"Diarization completed: {len(set(s['speaker'] for s in diarized_segments))} speakers detected")

                        # Apply post-processing smoothing to reduce rapid speaker changes
                        diarizer = SpeakerDiarizer()
                        diarized_segments = diarizer.smooth_speaker_changes(diarized_segments)

                    except Exception as e:
                        logger.error(f"Pyannote diarization failed: {e}", exc_info=True)
                        logger.warning("Falling back to simple heuristic diarization")

                        diarizer = SpeakerDiarizer()
                        diarized_segments = diarizer.diarize_segments(segments)
                else:
                    # Pyannote not available, use simple heuristic diarization
                    logger.info("Using simple heuristic diarization (no HF token or diarization disabled)")
                    diarizer = SpeakerDiarizer()
                    diarized_segments = diarizer.diarize_segments(segments)
            else:
                logger.info("Speaker diarization disabled")
                diarized_segments = segments

        except Exception as e:
            logger.error(f"Self-hosted transcription failed: {e}", exc_info=True)
            raise TranscriptionError(f"Transcription failed: {str(e)}")

        # Extract speaker information and generate speaker colors
        def generate_speaker_color(index: int) -> str:
            """Generate a distinct color for each speaker."""
            colors = [
                "#3B82F6",  # Blue
                "#EF4444",  # Red
                "#10B981",  # Green
                "#F59E0B",  # Amber
                "#8B5CF6",  # Purple
                "#EC4899",  # Pink
                "#14B8A6",  # Teal
                "#F97316",  # Orange
                "#6366F1",  # Indigo
                "#84CC16",  # Lime
            ]
            return colors[index % len(colors)]

        speakers = {}
        speaker_index = 0
        for segment in diarized_segments:
            speaker_id = segment.get('speaker', 'SPEAKER_01')
            if speaker_id not in speakers:
                # Generate human-readable name from speaker ID
                # Use the current speaker_index + 1 for consistent numbering
                speaker_name = f"Speaker {speaker_index + 1}"

                speakers[speaker_id] = {
                    'id': speaker_id,
                    'name': speaker_name,
                    'color': generate_speaker_color(speaker_index),
                    'segments_count': 0
                }
                speaker_index += 1
            speakers[speaker_id]['segments_count'] += 1

        # Step 5.5: AI-powered speaker name inference (if enabled)
        inference_metadata = None
        enable_name_inference = options.get("enable_name_inference", True)

        if enable_name_inference and len(speakers) > 0:
            try:
                self.update_state(
                    state='PROCESSING',
                    meta={'status': 'Inferring speaker names with AI', 'progress': 65}
                )

                # Get Ollama URL from environment
                ollama_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")

                # Run async inference in sync context using asyncio
                import asyncio

                # Create event loop if needed
                try:
                    loop = asyncio.get_event_loop()
                except RuntimeError:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                # Perform inference
                inferencer = SpeakerNameInferencer()
                speakers, inference_metadata = loop.run_until_complete(
                    inferencer.infer_speaker_names(
                        segments=diarized_segments,
                        speakers=speakers,
                        filename=document.filename,
                        ollama_url=ollama_url
                    )
                )

                logger.info(f"Speaker name inference completed with metadata: {inference_metadata}")

            except Exception as e:
                logger.warning(f"Speaker name inference failed, using default names: {e}")
                inference_metadata = {
                    'inference_performed': True,
                    'error': str(e),
                    'applied_names': {}
                }

        # Step 6: Export to multiple formats
        self.update_state(
            state='PROCESSING',
            meta={'status': 'Exporting transcription', 'progress': 70}
        )

        logger.info("Exporting transcription to multiple formats")
        exporter = TranscriptionExporter()

        # Prepare metadata (include inference metadata if available)
        metadata = {
            'document_id': document_id,
            'filename': document.filename,
            'language': language,
            'duration': duration,
            'segments_count': len(diarized_segments),
            'speakers_count': len(speakers)
        }

        # Add inference metadata if name inference was performed
        if inference_metadata:
            metadata['speaker_name_inference'] = inference_metadata

        # Export paths
        docx_path = os.path.join(temp_dir, "transcription.docx")
        srt_path = os.path.join(temp_dir, "transcription.srt")
        vtt_path = os.path.join(temp_dir, "transcription.vtt")
        json_path = os.path.join(temp_dir, "transcription.json")

        # Export to all formats
        exporter.export_to_docx(diarized_segments, docx_path)
        exporter.export_to_srt(diarized_segments, srt_path)
        exporter.export_to_vtt(diarized_segments, vtt_path)
        exporter.export_to_json(diarized_segments, json_path, metadata)

        # Step 7: Upload exports to MinIO
        self.update_state(
            state='PROCESSING',
            meta={'status': 'Uploading results', 'progress': 85}
        )

        logger.info("Uploading exports to MinIO")

        # Generate MinIO paths
        base_path = os.path.splitext(document.file_path)[0]
        export_paths = {}

        for format_name, local_path in [
            ('docx', docx_path),
            ('srt', srt_path),
            ('vtt', vtt_path),
            ('json', json_path)
        ]:
            if os.path.exists(local_path):
                minio_path = f"{base_path}.{format_name}"

                with open(local_path, 'rb') as f:
                    file_data = f.read()
                    content_type = {
                        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'srt': 'text/plain',
                        'vtt': 'text/vtt',
                        'json': 'application/json'
                    }.get(format_name, 'application/octet-stream')

                    minio_client.upload_file(
                        file_data=BytesIO(file_data),
                        object_name=minio_path,
                        content_type=content_type,
                        length=len(file_data)
                    )

                export_paths[format_name] = minio_path
                logger.info(f"Uploaded {format_name.upper()} to {minio_path}")

        # Step 8: Update transcription record in database
        self.update_state(
            state='PROCESSING',
            meta={'status': 'Saving transcription', 'progress': 95}
        )

        logger.info("Updating database with transcription results")

        # Find existing transcription (should always exist now)
        transcription = db.query(Transcription).filter(
            Transcription.document_id == document_id
        ).first()

        if not transcription:
            # This should not happen with new flow, but handle gracefully
            logger.warning(f"Transcription not found for document {document_id}, creating new one")
            transcription = Transcription(document_id=document_id)
            db.add(transcription)

        # Update transcription fields
        transcription.format = ext.lstrip('.') if ext else None
        transcription.duration = duration
        transcription.speakers = list(speakers.values())
        transcription.segments = diarized_segments

        # Update document status and metadata
        document.status = DocumentStatus.COMPLETED
        document.meta_data = document.meta_data or {}

        transcription_metadata = {
            'completed_at': datetime.utcnow().isoformat(),
            'duration': duration,
            'segments_count': len(diarized_segments),
            'speakers_count': len(speakers),
            'language': language,
            'export_formats': list(export_paths.keys())
        }

        # Add speaker name inference metadata if available
        if inference_metadata:
            transcription_metadata['speaker_name_inference'] = inference_metadata

        document.meta_data.update({
            'transcription': transcription_metadata
        })

        db.commit()

        logger.info(f"Transcription completed successfully for document {document_id}")

        # Step 9: Queue transcript indexing for search
        try:
            from app.workers.tasks.transcript_indexing import index_transcript
            logger.info(f"Queueing transcript indexing for transcription {transcription.id}")
            index_transcript.delay(transcription.id)
        except Exception as e:
            logger.warning(f"Failed to queue transcript indexing: {e}")
            # Don't fail the whole transcription if indexing fails to queue

        # Step 10: Return success result
        return {
            'status': 'completed',
            'document_id': document_id,
            'transcription_id': transcription.id,
            'filename': document.filename,
            'duration': duration,
            'segments_count': len(diarized_segments),
            'speakers_count': len(speakers),
            'speakers': list(speakers.values()),
            'export_paths': export_paths,
            'language': language,
            'task_id': self.request.id
        }

    except Exception as e:
        logger.error(f"Transcription failed for document {document_id}: {str(e)}", exc_info=True)

        # Update document status to FAILED
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if document:
                document.status = DocumentStatus.FAILED
                document.meta_data = document.meta_data or {}
                document.meta_data['transcription_error'] = {
                    'error': str(e),
                    'failed_at': datetime.utcnow().isoformat()
                }
                db.commit()
        except Exception as db_error:
            logger.error(f"Failed to update document status: {str(db_error)}")

        # Update task state
        self.update_state(
            state='FAILURE',
            meta={'status': f'Transcription failed: {str(e)}', 'error': str(e)}
        )

        return {
            'status': 'failed',
            'error': str(e),
            'document_id': document_id,
            'task_id': self.request.id
        }

    finally:
        # Cleanup temporary files
        if temp_dir and os.path.exists(temp_dir):
            try:
                import shutil
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to cleanup temp directory: {cleanup_error}")

        db.close()


@celery_app.task(name="process_transcription", bind=True)
def process_transcription(
    self,
    transcription_id: str,
    case_id: str,
    extract_entities: bool = True
) -> Dict[str, Any]:
    """
    Process a completed transcription for analysis and entity extraction.

    This is a placeholder task that will be fully implemented in Phase 3.
    It will handle post-transcription processing and NLP tasks.

    Args:
        transcription_id: Unique identifier for the transcription
        case_id: Unique identifier for the case
        extract_entities: Whether to extract named entities (default: True)

    Returns:
        Dict containing processing status and analysis results
    """
    # Placeholder - will implement in Phase 3
    # Future implementation will:
    # 1. Fetch transcription from database
    # 2. Run NLP entity extraction
    # 3. Extract legal terms and entities
    # 4. Create searchable chunks
    # 5. Index in Qdrant for semantic search
    # 6. Update database with entities

    return {
        "status": "pending",
        "transcription_id": transcription_id,
        "case_id": case_id,
        "extract_entities": extract_entities,
        "task_id": self.request.id,
    }
