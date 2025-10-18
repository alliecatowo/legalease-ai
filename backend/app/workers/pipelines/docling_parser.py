"""
Docling Document Parser

Extracts text, structure, and bounding boxes from various document formats using Docling library.
Supports PDF, DOCX, and other common legal document formats with OCR fallback.
"""

import logging
import io
from typing import Dict, Any, Optional, List
from pathlib import Path
import tempfile

logger = logging.getLogger(__name__)


class DoclingParser:
    """
    Document parser using Docling for text extraction and layout analysis.

    Features:
    - Multi-format support (PDF, DOCX, DOC, TXT)
    - Structure preservation (headers, paragraphs, tables)
    - OCR fallback for scanned documents
    - Bounding box extraction for visual highlighting
    - Metadata extraction
    """

    def __init__(self, use_ocr: bool = True, force_full_page_ocr: bool = False):
        """
        Initialize the Docling parser.

        Args:
            use_ocr: Whether to use OCR for scanned documents
            force_full_page_ocr: Force OCR on all pages (for scanned documents)
        """
        self.use_ocr = use_ocr
        self.force_full_page_ocr = force_full_page_ocr
        logger.info(f"Initialized DoclingParser (OCR: {use_ocr}, Force Full Page OCR: {force_full_page_ocr})")

    def parse(
        self,
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Parse a document and extract text, metadata, and bounding boxes.

        Args:
            file_content: Raw bytes of the document
            filename: Original filename (used to determine format)
            mime_type: MIME type of the document

        Returns:
            Dictionary containing:
                - text: Full extracted text
                - pages: List of page texts with bounding boxes
                - metadata: Document metadata
                - structure: Document structure information
                - bboxes: Bounding box data for visual highlighting
        """
        logger.info(f"Parsing document: {filename} (MIME: {mime_type})")

        # Determine file type
        file_ext = Path(filename).suffix.lower()

        try:
            # Route to appropriate parser based on file type
            if file_ext == '.pdf':
                return self._parse_pdf_with_docling(file_content, filename)
            elif file_ext in ['.docx', '.doc']:
                return self._parse_docx(file_content, filename)
            elif file_ext == '.txt':
                return self._parse_text(file_content, filename)
            else:
                logger.warning(f"Unsupported file type: {file_ext}, attempting generic parsing")
                return self._parse_generic(file_content, filename)

        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            # Fallback to PyMuPDF if Docling fails
            if file_ext == '.pdf':
                logger.warning("Falling back to PyMuPDF parser")
                return self._parse_pdf_fallback(file_content, filename)
            raise

    def _parse_pdf_with_docling(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse PDF document with Docling library for advanced layout analysis.
        Uses GPU acceleration when available and extracts detailed bounding boxes.

        Args:
            file_content: PDF file bytes
            filename: Original filename

        Returns:
            Parsed document data with bounding boxes
        """
        try:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.pipeline_options import PdfPipelineOptions, EasyOcrOptions
            from docling.datamodel.base_models import InputFormat
            import torch
        except ImportError as e:
            logger.error(f"Docling or torch not installed: {e}. Falling back to PyMuPDF")
            return self._parse_pdf_fallback(file_content, filename)

        try:
            # Detect GPU availability
            device = "cpu"
            if torch.cuda.is_available():
                device = "cuda"
                logger.info("GPU detected - enabling CUDA acceleration for Docling")
            else:
                logger.info("No GPU detected - using CPU for Docling")

            # Configure pipeline options with GPU support
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = self.use_ocr
            pipeline_options.do_table_structure = True
            pipeline_options.table_structure_options.do_cell_matching = True

            # Set device for GPU acceleration
            if hasattr(pipeline_options, 'accelerator_options'):
                pipeline_options.accelerator_options.device = device

            # Configure OCR options for scanned documents
            if self.use_ocr:
                ocr_options = EasyOcrOptions(force_full_page_ocr=self.force_full_page_ocr)
                # Enable GPU for OCR if available
                if device == "cuda" and hasattr(ocr_options, 'use_gpu'):
                    ocr_options.use_gpu = True
                pipeline_options.ocr_options = ocr_options

            # Create document converter
            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )

            # Write to temporary file (Docling needs file path)
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name

            try:
                # Estimate processing time based on file size
                file_size_mb = len(file_content) / (1024 * 1024)
                # Rough estimate: 1-2 seconds per MB on GPU, 5-10 seconds on CPU
                est_time_sec = file_size_mb * (1.5 if device == "cuda" else 7)
                logger.info(
                    f"Converting {file_size_mb:.1f}MB PDF with {device.upper()} "
                    f"(estimated ~{int(est_time_sec/60)} minutes)..."
                )

                # Convert document (this is the long-running operation with no progress callbacks)
                result = converter.convert(tmp_path)
                doc = result.document

                # Extract pages with bounding boxes using doc.iterate_items()
                pages_dict = {}  # page_num -> {"text": [], "items": []}
                full_text = []

                # Use iterate_items() to get proper tuple structure
                if hasattr(doc, 'iterate_items'):
                    logger.info("Using doc.iterate_items() for bbox extraction")

                    for item_tuple in doc.iterate_items():
                        # Items are tuples: (ContentObject, page_number)
                        if not isinstance(item_tuple, tuple) or len(item_tuple) < 2:
                            continue

                        content_obj = item_tuple[0]
                        page_num = item_tuple[1]

                        # Initialize page if not seen
                        if page_num not in pages_dict:
                            pages_dict[page_num] = {
                                "text": [],
                                "items": []
                            }

                        # Extract text
                        item_text = getattr(content_obj, 'text', '')
                        if item_text:
                            pages_dict[page_num]["text"].append(item_text)

                        # Extract bounding box from prov structure
                        prov = getattr(content_obj, 'prov', None)
                        bbox = None

                        if prov and len(prov) > 0:
                            bbox_obj = getattr(prov[0], 'bbox', None)
                            if bbox_obj:
                                bbox = {
                                    "l": getattr(bbox_obj, 'l', 0),
                                    "t": getattr(bbox_obj, 't', 0),
                                    "r": getattr(bbox_obj, 'r', 0),
                                    "b": getattr(bbox_obj, 'b', 0)
                                }

                        # Create item entry
                        item_data = {
                            "text": item_text,
                            "type": content_obj.__class__.__name__,
                        }

                        if bbox:
                            item_data["bbox"] = bbox

                        pages_dict[page_num]["items"].append(item_data)

                else:
                    # Fallback to old method if iterate_items not available
                    logger.warning("doc.iterate_items() not available, using page.items fallback")
                    for page_num, page in enumerate(doc.pages):
                        if page_num not in pages_dict:
                            pages_dict[page_num] = {
                                "text": [],
                                "items": []
                            }

                        for item in page.items:
                            item_text = getattr(item, 'text', '')
                            if item_text:
                                pages_dict[page_num]["text"].append(item_text)

                            # Try to get bbox
                            bbox = None
                            if hasattr(item, 'bbox') and item.bbox:
                                bbox_obj = item.bbox
                                bbox = {
                                    "l": getattr(bbox_obj, 'l', 0),
                                    "t": getattr(bbox_obj, 't', 0),
                                    "r": getattr(bbox_obj, 'r', 0),
                                    "b": getattr(bbox_obj, 'b', 0)
                                }

                            item_data = {
                                "text": item_text,
                                "type": item.__class__.__name__,
                            }

                            if bbox:
                                item_data["bbox"] = bbox

                            pages_dict[page_num]["items"].append(item_data)

                # Build pages list in order
                pages = []
                for page_num in sorted(pages_dict.keys()):
                    page_data = pages_dict[page_num]
                    page_text = "\n".join(page_data["text"])

                    pages.append({
                        "page_number": page_num,
                        "text": page_text,
                        "items": page_data["items"],
                    })
                    full_text.append(page_text)

                # Extract metadata
                metadata = {
                    "page_count": len(pages),
                    "title": doc.metadata.get("title", "") if hasattr(doc, 'metadata') else "",
                    "author": doc.metadata.get("author", "") if hasattr(doc, 'metadata') else "",
                    "device": device,  # Track whether GPU was used
                }

                # Count total items with bboxes for logging
                total_items = sum(len(p["items"]) for p in pages)
                items_with_bbox = sum(
                    1 for p in pages for item in p["items"] if "bbox" in item
                )

                result = {
                    "text": "\n\n".join(full_text),
                    "pages": pages,
                    "metadata": metadata,
                }

                logger.info(
                    f"Successfully parsed PDF with Docling ({device}): "
                    f"{len(pages)} pages, {total_items} items, "
                    f"{items_with_bbox} with bboxes "
                    f"({items_with_bbox/total_items*100:.1f}% coverage)"
                )

                # Force GPU memory cleanup after Docling processing
                # This prevents GPU OOM when embedding pipeline loads its models
                if device == "cuda":
                    try:
                        import torch
                        import gc

                        # Clear PyTorch cache
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                            torch.cuda.synchronize()
                            logger.info("Cleared GPU memory after Docling processing")

                        # Force garbage collection to release Docling models
                        gc.collect()
                    except Exception as e:
                        logger.warning(f"Failed to clear GPU memory: {e}")

                return result

            finally:
                # Clean up temporary file
                import os
                try:
                    os.unlink(tmp_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file: {e}")

                # Clear GPU cache if used
                if device == "cuda":
                    try:
                        torch.cuda.empty_cache()
                        logger.debug("Cleared CUDA cache")
                    except Exception as e:
                        logger.debug(f"Failed to clear CUDA cache: {e}")

        except Exception as e:
            logger.error(f"Error parsing PDF with Docling: {e}", exc_info=True)
            logger.warning("Falling back to PyMuPDF")
            return self._parse_pdf_fallback(file_content, filename)

    def _parse_pdf_fallback(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Fallback PDF parser using PyMuPDF with basic bbox extraction.

        Args:
            file_content: PDF file bytes
            filename: Original filename

        Returns:
            Parsed document data
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF (fitz) not installed. Install with: pip install pymupdf")
            raise

        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=file_content, filetype="pdf")

            pages = []
            full_text = []
            all_bboxes = []
            metadata = {}

            # Extract metadata
            metadata = {
                "page_count": pdf_document.page_count,
                "title": pdf_document.metadata.get("title", ""),
                "author": pdf_document.metadata.get("author", ""),
                "subject": pdf_document.metadata.get("subject", ""),
                "creator": pdf_document.metadata.get("creator", ""),
            }

            # Extract text and bboxes from each page
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]

                # Get text with word-level bounding boxes
                words = page.get_text("words")  # Returns list of (x0, y0, x1, y1, word, block_no, line_no, word_no)

                page_text = page.get_text()
                page_bboxes = []

                # Extract bounding boxes for each word
                for word_data in words:
                    if len(word_data) >= 5:
                        x0, y0, x1, y1, word = word_data[:5]
                        bbox_data = {
                            "x0": x0,
                            "y0": y0,
                            "x1": x1,
                            "y1": y1,
                            "page": page_num + 1,
                            "text": word,
                            "type": "word"
                        }
                        page_bboxes.append(bbox_data)
                        all_bboxes.append(bbox_data)

                # If page is mostly empty and OCR is enabled, try OCR
                if self.use_ocr and len(page_text.strip()) < 50:
                    logger.info(f"Page {page_num + 1} has little text, attempting OCR")
                    page_text = self._ocr_page(page)

                pages.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "char_count": len(page_text),
                    "bboxes": page_bboxes,
                })
                full_text.append(page_text)

            pdf_document.close()

            result = {
                "text": "\n\n".join(full_text),
                "pages": pages,
                "metadata": metadata,
                "structure": {
                    "type": "pdf",
                    "page_count": len(pages),
                },
                "bboxes": all_bboxes,
            }

            logger.info(f"Successfully parsed PDF (PyMuPDF): {len(pages)} pages, {len(all_bboxes)} bboxes")
            return result

        except Exception as e:
            logger.error(f"Error in PDF fallback parser: {e}")
            raise

    def _parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse DOCX document.

        Args:
            file_content: DOCX file bytes
            filename: Original filename

        Returns:
            Parsed document data
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise

        try:
            # Open DOCX from bytes
            doc = Document(io.BytesIO(file_content))

            paragraphs = []
            full_text = []

            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                    })
                    full_text.append(para.text)

            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "index": table_idx,
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "data": table_data,
                })

            # Extract metadata
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            }

            result = {
                "text": "\n\n".join(full_text),
                "pages": [],  # DOCX doesn't have explicit pages
                "metadata": metadata,
                "structure": {
                    "type": "docx",
                    "paragraphs": paragraphs,
                    "tables": tables,
                },
                "bboxes": [],  # DOCX doesn't have bounding boxes
            }

            logger.info(f"Successfully parsed DOCX: {len(paragraphs)} paragraphs, {len(tables)} tables")
            return result

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise

    def _parse_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse plain text document.

        Args:
            file_content: Text file bytes
            filename: Original filename

        Returns:
            Parsed document data
        """
        try:
            # Try UTF-8 first, then fall back to latin-1
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')

            # Split into lines
            lines = text.split('\n')

            result = {
                "text": text,
                "pages": [],
                "metadata": {
                    "line_count": len(lines),
                    "char_count": len(text),
                },
                "structure": {
                    "type": "text",
                    "lines": len(lines),
                },
                "bboxes": [],  # Text files don't have bounding boxes
            }

            logger.info(f"Successfully parsed text file: {len(lines)} lines")
            return result

        except Exception as e:
            logger.error(f"Error parsing text file: {e}")
            raise

    def _parse_generic(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Generic fallback parser for unsupported formats.

        Args:
            file_content: File bytes
            filename: Original filename

        Returns:
            Parsed document data (best effort)
        """
        logger.warning(f"Using generic parser for {filename}")

        try:
            # Try to decode as text
            text = file_content.decode('utf-8', errors='ignore')

            return {
                "text": text,
                "pages": [],
                "metadata": {
                    "char_count": len(text),
                    "format": "unknown",
                },
                "structure": {
                    "type": "generic",
                },
                "bboxes": [],
            }
        except Exception as e:
            logger.error(f"Error in generic parser: {e}")
            raise

    def _ocr_page(self, page) -> str:
        """
        Perform OCR on a PDF page.

        Args:
            page: PyMuPDF page object

        Returns:
            OCR-extracted text
        """
        try:
            import fitz  # PyMuPDF
            import pytesseract
            from PIL import Image
        except ImportError:
            logger.warning("OCR dependencies not installed (pytesseract, Pillow, PyMuPDF)")
            return ""

        try:
            # Render page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Perform OCR
            text = pytesseract.image_to_string(img)
            logger.info(f"OCR extracted {len(text)} characters")

            return text

        except Exception as e:
            logger.error(f"Error performing OCR: {e}")
            return ""

    def validate_document(self, file_content: bytes, filename: str) -> bool:
        """
        Validate if a document can be parsed.

        Args:
            file_content: File bytes
            filename: Original filename

        Returns:
            True if document can be parsed
        """
        file_ext = Path(filename).suffix.lower()
        supported_formats = ['.pdf', '.docx', '.doc', '.txt']

        return file_ext in supported_formats
