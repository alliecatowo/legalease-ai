"""Transcription service for managing audio/video transcription operations."""

import io
import json
import logging
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional
from uuid import UUID

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile
from minio.error import S3Error
from sqlalchemy.orm import Session, joinedload

from app.core.minio_client import minio_client
from app.models.case import Case
from app.models.document import DocumentStatus
from app.models.transcription import TranscriptSegment, Transcription

logger = logging.getLogger(__name__)


class TranscriptionService:
    """Service class for transcription operations."""

    # Supported audio/video formats
    SUPPORTED_FORMATS = {
        # Audio formats
        "audio/mpeg",
        "audio/mp3",
        "audio/mp4",  # M4A files (MPEG-4 audio)
        "audio/wav",
        "audio/x-wav",
        "audio/wave",
        "audio/vnd.wave",  # Alternative MIME type for WAV files
        "audio/aac",
        "audio/m4a",
        "audio/x-m4a",
        "audio/flac",
        "audio/ogg",
        "audio/webm",
        # Video formats
        "video/mp4",
        "video/mpeg",
        "video/quicktime",
        "video/x-msvideo",
        "video/webm",
        "video/x-matroska",
    }

    @staticmethod
    async def upload_audio_for_transcription(
        case_id: int,
        file: UploadFile,
        db: Session,
        options: Optional["TranscriptionOptions"] = None,
        team_id: Optional[UUID] = None,
    ) -> Transcription:
        """
        Upload an audio/video file for transcription.

        Args:
            case_id: ID of the case
            file: Uploaded audio/video file
            db: Database session
            options: Optional transcription configuration options

        Returns:
            Transcription: Created transcription record

        Raises:
            HTTPException: If validation fails or upload fails
        """
        # Validate case exists
        query = db.query(Case).filter(Case.id == case_id)
        if team_id:
            query = query.filter(Case.team_id == team_id)
        case = query.first()
        if not case:
            logger.error(f"Case {case_id} not found")
            raise HTTPException(status_code=404, detail=f"Case {case_id} not found")

        # Validate file type
        if file.content_type not in TranscriptionService.SUPPORTED_FORMATS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format: {file.content_type}. "
                f"Supported formats: audio (mp3, wav, aac, m4a, flac, ogg, webm) "
                f"and video (mp4, mpeg, mov, avi, webm, mkv)",
            )

        # Generate unique file identifier
        file_uuid = str(uuid.uuid4())
        filename = file.filename or "unknown"

        # Create MinIO path: cases/{case_id}/transcripts/{uuid}_{filename}
        object_name = f"cases/{case_id}/transcripts/{file_uuid}_{filename}"

        # Read file content and get size
        file_content = await file.read()
        file_size = len(file_content)

        # Upload to MinIO
        try:
            minio_client.upload_file(
                file_data=io.BytesIO(file_content),
                object_name=object_name,
                content_type=file.content_type,
                length=file_size,
            )
            logger.info(f"Uploaded file to MinIO: {object_name}")
        except S3Error as e:
            logger.error(f"Failed to upload file to MinIO: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to upload file to storage: {str(e)}",
            )

        # Extract format from content type (e.g., "audio/mp3" -> "mp3")
        file_format = file.content_type.split("/")[-1] if file.content_type else None

        # Create Transcription record directly
        transcription = Transcription(
            case_id=case_id,
            filename=filename,
            file_path=object_name,
            mime_type=file.content_type,
            size=file_size,
            status=DocumentStatus.PENDING,
            uploaded_at=datetime.utcnow(),
            format=file_format,
            segments=[],
            speakers=[],
        )

        db.add(transcription)
        db.commit()
        db.refresh(transcription)

        # Queue transcription task with options
        from app.workers.tasks.transcription import transcribe_audio

        options_dict = options.model_dump() if options else {}
        transcribe_audio.delay(transcription.id, options=options_dict)

        logger.info(f"Created transcription {transcription.id} and queued task with options: {options_dict}")

        return transcription

    @staticmethod
    def get_transcription(
        transcription_id: int,
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> Transcription:
        """
        Get a transcription by ID.

        Args:
            transcription_id: ID of the transcription
            db: Database session

        Returns:
            Transcription: Transcription record

        Raises:
            HTTPException: If transcription not found
        """
        query = (
            db.query(Transcription)
            .options(joinedload(Transcription.case))
            .join(Case, Transcription.case_id == Case.id)
            .filter(Transcription.id == transcription_id)
        )
        if team_id:
            query = query.filter(Case.team_id == team_id)

        transcription = query.first()
        if not transcription:
            logger.error(f"Transcription {transcription_id} not found")
            raise HTTPException(
                status_code=404, detail=f"Transcription {transcription_id} not found"
            )

        return transcription

    @staticmethod
    def list_case_transcriptions(
        case_id: int,
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> List[Dict[str, Any]]:
        """
        List all transcriptions for a case.

        Args:
            case_id: ID of the case
            db: Database session

        Returns:
            List[Dict]: List of transcription summaries

        Raises:
            HTTPException: If case not found
        """
        # Validate case exists
        query = db.query(Case).filter(Case.id == case_id)
        if team_id:
            query = query.filter(Case.team_id == team_id)
        case = query.first()
        if not case:
            logger.error(f"Case {case_id} not found")
            raise HTTPException(status_code=404, detail=f"Case {case_id} not found")

        # Get all transcriptions for the case directly
        transcriptions = (
            db.query(Transcription)
            .filter(Transcription.case_id == case_id)
            .all()
        )

        logger.info(f"Found {len(transcriptions)} transcriptions for case {case_id}")

        # Build summary list
        result = []
        for trans in transcriptions:
            result.append(
                {
                    "id": trans.id,
                    "case_id": trans.case_id,
                    "filename": trans.filename,
                    "format": trans.format,
                    "duration": trans.duration,
                    "segment_count": len(trans.segments) if trans.segments else 0,
                    "speaker_count": len(trans.speakers) if trans.speakers else 0,
                    "status": trans.status.value if trans.status else "unknown",
                    "created_at": trans.created_at,
                    "uploaded_at": trans.uploaded_at,
                }
            )

        return result

    @staticmethod
    def download_audio(
        transcription_id: int,
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> tuple[bytes, str, str]:
        """
        Download the original audio/video file from MinIO.

        Args:
            transcription_id: ID of the transcription
            db: Database session

        Returns:
            tuple: (file_content, filename, content_type)

        Raises:
            HTTPException: If transcription not found or download fails
        """
        # Get transcription from database
        transcription = TranscriptionService.get_transcription(
            transcription_id,
            db,
            team_id=team_id,
        )

        try:
            # Download from MinIO
            logger.info(f"Downloading audio for transcription {transcription_id} from MinIO: {transcription.file_path}")
            content = minio_client.download_file(transcription.file_path)

            return content, transcription.filename, transcription.mime_type or "application/octet-stream"

        except S3Error as e:
            logger.error(f"MinIO error downloading audio for transcription {transcription_id}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to download audio from storage: {str(e)}",
            )
        except Exception as e:
            logger.error(f"Error downloading audio for transcription {transcription_id}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to download audio: {str(e)}",
            )

    @staticmethod
    def get_transcription_details(
        transcription_id: int,
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> Dict[str, Any]:
        """
        Get detailed transcription information.

        Args:
            transcription_id: ID of the transcription
            db: Database session

        Returns:
            Dict: Detailed transcription data

        Raises:
            HTTPException: If transcription not found
        """
        transcription = TranscriptionService.get_transcription(
            transcription_id,
            db,
            team_id=team_id,
        )

        # Get key moment metadata for segments
        key_moment_metadata = (
            db.query(TranscriptSegment)
            .filter(TranscriptSegment.transcript_id == transcription_id)
            .all()
        )

        # Build a dict of segment_id -> is_key_moment for quick lookup
        key_moments_map = {meta.segment_id: meta.is_key_moment for meta in key_moment_metadata}

        # Merge key moment status into segments
        segments_with_metadata = []
        for segment in (transcription.segments or []):
            segment_copy = segment.copy()
            segment_id = segment.get('id')
            # Add isKeyMoment field (camelCase for frontend)
            segment_copy['isKeyMoment'] = key_moments_map.get(segment_id, False)
            segments_with_metadata.append(segment_copy)

        return {
            "id": transcription.id,
            "case_id": transcription.case_id,
            "filename": transcription.filename,
            "format": transcription.format,
            "duration": transcription.duration,
            "speakers": transcription.speakers,
            "segments": segments_with_metadata,
            "status": transcription.status.value if transcription.status else "unknown",
            "created_at": transcription.created_at,
            "uploaded_at": transcription.uploaded_at,
            "audio_url": f"/api/v1/transcriptions/{transcription.id}/audio",
        }

    @staticmethod
    def export_as_json(transcription: Transcription, db: Session) -> bytes:
        """Export transcription as JSON."""
        data = {
            "transcription_id": transcription.id,
            "case_id": transcription.case_id,
            "filename": transcription.filename,
            "format": transcription.format,
            "duration": transcription.duration,
            "speakers": transcription.speakers,
            "segments": transcription.segments,
            "created_at": transcription.created_at.isoformat(),
            "uploaded_at": transcription.uploaded_at.isoformat() if transcription.uploaded_at else None,
        }

        return json.dumps(data, indent=2).encode("utf-8")

    @staticmethod
    def export_as_docx(transcription: Transcription, db: Session) -> bytes:
        """Export transcription as DOCX with formatting."""
        doc = DocxDocument()

        # Add title
        title = doc.add_heading(
            f"Transcription: {transcription.filename}", 0
        )

        # Add metadata
        doc.add_paragraph(f"Transcription ID: {transcription.id}")
        doc.add_paragraph(f"Case ID: {transcription.case_id}")
        doc.add_paragraph(f"Format: {transcription.format or 'Unknown'}")
        if transcription.duration:
            minutes = int(transcription.duration // 60)
            seconds = int(transcription.duration % 60)
            doc.add_paragraph(f"Duration: {minutes}m {seconds}s")
        doc.add_paragraph(f"Created: {transcription.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        if transcription.uploaded_at:
            doc.add_paragraph(f"Uploaded: {transcription.uploaded_at.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        # Add speakers section if available
        if transcription.speakers:
            doc.add_heading("Speakers", level=1)
            for speaker in transcription.speakers:
                speaker_text = f"• {speaker.get('speaker_id', 'Unknown')}"
                if speaker.get("label"):
                    speaker_text += f" - {speaker['label']}"
                doc.add_paragraph(speaker_text)
            doc.add_paragraph()

        # Add transcription content
        doc.add_heading("Transcription", level=1)

        segments = transcription.segments or []
        for segment in segments:
            # Format timestamp
            start_time = TranscriptionService._format_timestamp(segment.get("start", 0))
            end_time = TranscriptionService._format_timestamp(segment.get("end", 0))
            timestamp = f"[{start_time} - {end_time}]"

            # Create paragraph with speaker and timestamp
            p = doc.add_paragraph()
            if segment.get("speaker"):
                p.add_run(f"{segment['speaker']}: ").bold = True
            p.add_run(f"{timestamp} ")
            p.add_run(segment.get("text", ""))
            p.add_run("\n")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def export_as_srt(transcription: Transcription) -> bytes:
        """Export transcription as SRT subtitle format."""
        srt_content = []
        segments = transcription.segments or []

        for idx, segment in enumerate(segments, 1):
            start = TranscriptionService._format_srt_timestamp(segment.get("start", 0))
            end = TranscriptionService._format_srt_timestamp(segment.get("end", 0))
            text = segment.get("text", "")

            # Add speaker label if available
            if segment.get("speaker"):
                text = f"[{segment['speaker']}] {text}"

            srt_content.append(f"{idx}")
            srt_content.append(f"{start} --> {end}")
            srt_content.append(text)
            srt_content.append("")  # Empty line between subtitles

        return "\n".join(srt_content).encode("utf-8")

    @staticmethod
    def export_as_vtt(transcription: Transcription) -> bytes:
        """Export transcription as VTT subtitle format."""
        vtt_content = ["WEBVTT", ""]
        segments = transcription.segments or []

        for segment in segments:
            start = TranscriptionService._format_vtt_timestamp(segment.get("start", 0))
            end = TranscriptionService._format_vtt_timestamp(segment.get("end", 0))
            text = segment.get("text", "")

            # Add speaker label if available
            if segment.get("speaker"):
                text = f"<v {segment['speaker']}>{text}"

            vtt_content.append(f"{start} --> {end}")
            vtt_content.append(text)
            vtt_content.append("")  # Empty line between subtitles

        return "\n".join(vtt_content).encode("utf-8")

    @staticmethod
    def export_as_txt(transcription: Transcription) -> bytes:
        """Export transcription as plain text."""
        txt_content = []
        segments = transcription.segments or []

        for segment in segments:
            start_time = TranscriptionService._format_timestamp(segment.get("start", 0))
            text = segment.get("text", "")

            # Add speaker and timestamp
            line = f"[{start_time}]"
            if segment.get("speaker"):
                line += f" {segment['speaker']}:"
            line += f" {text}"

            txt_content.append(line)

        return "\n".join(txt_content).encode("utf-8")

    @staticmethod
    def delete_transcription(
        transcription_id: int,
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> Transcription:
        """
        Delete a transcription.

        Args:
            transcription_id: ID of the transcription
            db: Database session

        Returns:
            Transcription: Deleted transcription record

        Raises:
            HTTPException: If transcription not found or deletion fails
        """
        transcription = TranscriptionService.get_transcription(
            transcription_id,
            db,
            team_id=team_id,
        )

        try:
            # Delete from database (cascade will handle related records)
            db.delete(transcription)
            db.commit()

            logger.info(f"Successfully deleted transcription {transcription_id}")
            return transcription

        except Exception as e:
            logger.error(f"Error deleting transcription {transcription_id}: {str(e)}")
            db.rollback()
            raise HTTPException(
                status_code=500,
                detail=f"Failed to delete transcription: {str(e)}",
            )

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """Format seconds as HH:MM:SS."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

    @staticmethod
    def _format_srt_timestamp(seconds: float) -> str:
        """Format seconds as SRT timestamp (HH:MM:SS,mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    @staticmethod
    def _format_vtt_timestamp(seconds: float) -> str:
        """Format seconds as VTT timestamp (HH:MM:SS.mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"

    @staticmethod
    def toggle_key_moment(
        transcription_id: int,
        segment_id: str,
        is_key_moment: bool,
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> Dict[str, Any]:
        """
        Toggle key moment status for a transcript segment.

        Args:
            transcription_id: ID of the transcription
            segment_id: UUID of the segment (from JSON segments)
            is_key_moment: New key moment status
            db: Database session

        Returns:
            Dict: Updated segment metadata

        Raises:
            HTTPException: If transcription or segment not found
        """
        # Verify transcription exists
        transcription = TranscriptionService.get_transcription(
            transcription_id,
            db,
            team_id=team_id,
        )

        # Verify segment exists in the transcription's JSON segments
        segment_exists = False
        if transcription.segments:
            for seg in transcription.segments:
                if seg.get("id") == segment_id:
                    segment_exists = True
                    break

        if not segment_exists:
            logger.error(
                f"Segment {segment_id} not found in transcription {transcription_id}"
            )
            raise HTTPException(
                status_code=404,
                detail=f"Segment {segment_id} not found in transcription {transcription_id}",
            )

        # Check if segment metadata already exists
        segment_metadata = (
            db.query(TranscriptSegment)
            .filter(
                TranscriptSegment.transcript_id == transcription_id,
                TranscriptSegment.segment_id == segment_id,
            )
            .first()
        )

        if segment_metadata:
            # Update existing metadata
            segment_metadata.is_key_moment = is_key_moment
            segment_metadata.updated_at = datetime.utcnow()
            db.commit()
            db.refresh(segment_metadata)
        else:
            # Create new metadata
            segment_metadata = TranscriptSegment(
                transcript_id=transcription_id,
                segment_id=segment_id,
                is_key_moment=is_key_moment,
            )
            db.add(segment_metadata)
            db.commit()
            db.refresh(segment_metadata)

        logger.info(
            f"Toggled key moment for segment {segment_id} in transcription {transcription_id}: {is_key_moment}"
        )

        return {
            "segment_id": segment_metadata.segment_id,
            "is_key_moment": segment_metadata.is_key_moment,
            "updated_at": segment_metadata.updated_at.isoformat(),
        }

    @staticmethod
    def get_key_moments(
        transcription_id: int,
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> Dict[str, Any]:
        """
        Get all key moments for a transcription.

        Args:
            transcription_id: ID of the transcription
            db: Database session

        Returns:
            Dict: Key moments data with full segment information

        Raises:
            HTTPException: If transcription not found
        """
        # Verify transcription exists
        transcription = TranscriptionService.get_transcription(
            transcription_id,
            db,
            team_id=team_id,
        )

        # Get all key moment segment metadata
        key_moment_metadata = (
            db.query(TranscriptSegment)
            .filter(
                TranscriptSegment.transcript_id == transcription_id,
                TranscriptSegment.is_key_moment == True,
            )
            .all()
        )

        # Build a set of key moment segment IDs for efficient lookup
        key_moment_ids = {meta.segment_id for meta in key_moment_metadata}

        # Filter segments from JSON to get full segment data
        key_moments = []
        if transcription.segments:
            for seg in transcription.segments:
                if seg.get("id") in key_moment_ids:
                    key_moments.append(
                        {
                            "segment_id": seg.get("id"),
                            "text": seg.get("text", ""),
                            "speaker": seg.get("speaker"),
                            "start_time": seg.get("start"),
                            "end_time": seg.get("end"),
                            "confidence": seg.get("confidence"),
                        }
                    )

        # Sort by start time
        key_moments.sort(key=lambda x: x.get("start_time", 0))

        logger.info(
            f"Retrieved {len(key_moments)} key moments for transcription {transcription_id}"
        )

        return {
            "transcription_id": transcription_id,
            "key_moments": key_moments,
            "total": len(key_moments),
        }

    @staticmethod
    def update_speaker(
        transcription_id: int,
        speaker_id: str,
        name: str,
        role: Optional[str],
        db: Session,
        team_id: Optional[UUID] = None,
    ) -> Dict[str, Any]:
        """
        Update speaker information in a transcription.

        Args:
            transcription_id: ID of the transcription
            speaker_id: ID of the speaker to update
            name: New speaker name
            role: Optional speaker role
            db: Database session

        Returns:
            Dict: Updated speaker information

        Raises:
            HTTPException: If transcription or speaker not found
        """
        # Verify transcription exists
        transcription = TranscriptionService.get_transcription(
            transcription_id,
            db,
            team_id=team_id,
        )

        # Verify speakers array exists
        if not transcription.speakers or not isinstance(transcription.speakers, list):
            logger.error(
                f"No speakers found in transcription {transcription_id}"
            )
            raise HTTPException(
                status_code=404,
                detail=f"No speakers found in transcription {transcription_id}",
            )

        # Find the speaker by ID
        speaker_found = False
        updated_speaker = None
        for speaker in transcription.speakers:
            if speaker.get("speaker_id") == speaker_id:
                # Update speaker information
                speaker["name"] = name
                if role is not None:
                    speaker["role"] = role
                speaker_found = True
                updated_speaker = speaker.copy()
                break

        if not speaker_found:
            logger.error(
                f"Speaker {speaker_id} not found in transcription {transcription_id}"
            )
            raise HTTPException(
                status_code=404,
                detail=f"Speaker {speaker_id} not found in transcription {transcription_id}",
            )

        # Mark the speakers column as modified for SQLAlchemy to detect the change
        from sqlalchemy.orm.attributes import flag_modified
        flag_modified(transcription, "speakers")

        # Commit the changes
        db.commit()
        db.refresh(transcription)

        logger.info(
            f"Updated speaker {speaker_id} in transcription {transcription_id}: name='{name}', role='{role}'"
        )

        return {
            "speaker_id": updated_speaker.get("speaker_id"),
            "name": updated_speaker.get("name"),
            "role": updated_speaker.get("role"),
            "color": updated_speaker.get("color"),
        }
