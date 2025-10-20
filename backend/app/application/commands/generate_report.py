"""
GenerateReportCommand - Generates dossier files from research results.

This command creates DOCX and/or PDF files from a completed ResearchRun,
uploads them to MinIO, and updates the dossier record.
"""

import logging
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import List, Optional
from uuid import UUID

from app.domain.research.repositories.research_repository import (
    DossierRepository,
    ResearchRunRepository,
)


logger = logging.getLogger(__name__)


@dataclass
class GenerateReportCommand:
    """
    Command to generate dossier reports.

    Attributes:
        research_run_id: UUID of the research run
        format: Report format - "DOCX", "PDF", or "BOTH"
        include_citations: Whether to include citations appendix
        template_name: Optional custom template name
    """

    research_run_id: UUID
    format: str = "BOTH"
    include_citations: bool = True
    template_name: Optional[str] = None


@dataclass
class GenerateReportResult:
    """
    Result of report generation.

    Attributes:
        success: Whether generation succeeded
        file_paths: List of generated file paths (MinIO URLs)
        message: Human-readable status message
        error: Error message if failed
        metadata: Additional result metadata
    """

    success: bool
    file_paths: List[str] = None
    message: str = ""
    error: Optional[str] = None
    metadata: Optional[dict] = None

    def __post_init__(self):
        if self.file_paths is None:
            self.file_paths = []


class GenerateReportCommandHandler:
    """
    Handler for GenerateReportCommand.

    Orchestrates:
    1. Retrieving dossier from database
    2. Rendering to DOCX using template
    3. Converting to PDF if needed
    4. Uploading to MinIO
    5. Updating dossier record
    """

    def __init__(
        self,
        dossier_repository: DossierRepository,
        research_repository: ResearchRunRepository,
        storage_service: any,  # MinIO client
        template_path: str = "/app/templates/dossier",
    ):
        """
        Initialize handler with dependencies.

        Args:
            dossier_repository: Repository for dossier persistence
            research_repository: Repository for research run access
            storage_service: MinIO client for file storage
            template_path: Path to report templates
        """
        self.dossier_repo = dossier_repository
        self.research_repo = research_repository
        self.storage = storage_service
        self.template_path = Path(template_path)

    async def handle(self, command: GenerateReportCommand) -> GenerateReportResult:
        """
        Handle the GenerateReportCommand.

        Args:
            command: The command to execute

        Returns:
            GenerateReportResult with file paths

        Process:
            1. Get research run and dossier from database
            2. Validate research is completed
            3. Render DOCX from template
            4. Convert to PDF if requested
            5. Upload to MinIO
            6. Update dossier with file paths
            7. Return result
        """
        logger.info(f"Generating report for research run {command.research_run_id}")

        try:
            # Get research run
            research_run = await self.research_repo.get_by_id(command.research_run_id)
            if not research_run:
                return GenerateReportResult(
                    success=False,
                    message=f"Research run {command.research_run_id} not found",
                    error="Research run not found",
                )

            # Verify research is completed
            if not research_run.is_completed():
                return GenerateReportResult(
                    success=False,
                    message="Research run is not completed yet",
                    error="Research not completed",
                )

            # Get dossier
            dossier = await self.dossier_repo.get_by_research_run(
                command.research_run_id
            )
            if not dossier:
                return GenerateReportResult(
                    success=False,
                    message=f"Dossier for research run {command.research_run_id} not found",
                    error="Dossier not found",
                )

            # Generate file paths
            file_paths = []
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            base_filename = f"dossier_{command.research_run_id}_{timestamp}"

            # Generate DOCX
            if command.format in ("DOCX", "BOTH"):
                docx_path = await self._generate_docx(
                    dossier, base_filename, command.include_citations, command.template_name
                )
                file_paths.append(docx_path)
                logger.info(f"Generated DOCX: {docx_path}")

            # Generate PDF
            if command.format in ("PDF", "BOTH"):
                pdf_path = await self._generate_pdf(
                    dossier, base_filename, command.include_citations, command.template_name
                )
                file_paths.append(pdf_path)
                logger.info(f"Generated PDF: {pdf_path}")

            # Update dossier metadata with file paths
            dossier.metadata["generated_files"] = file_paths
            dossier.metadata["last_generated"] = datetime.utcnow().isoformat()
            await self.dossier_repo.save(dossier)

            logger.info(
                f"Successfully generated {len(file_paths)} report(s) for research run {command.research_run_id}"
            )

            return GenerateReportResult(
                success=True,
                file_paths=file_paths,
                message=f"Successfully generated {len(file_paths)} report file(s)",
                metadata={
                    "research_run_id": str(command.research_run_id),
                    "format": command.format,
                    "word_count": dossier.get_word_count(),
                },
            )

        except Exception as e:
            logger.error(
                f"Failed to generate report for research run {command.research_run_id}: {e}",
                exc_info=True,
            )
            return GenerateReportResult(
                success=False,
                message="Failed to generate report",
                error=str(e),
            )

    async def _generate_docx(
        self,
        dossier,
        base_filename: str,
        include_citations: bool,
        template_name: Optional[str],
    ) -> str:
        """
        Generate DOCX file from dossier.

        Args:
            dossier: Dossier entity
            base_filename: Base filename without extension
            include_citations: Whether to include citations
            template_name: Optional template name

        Returns:
            MinIO path to uploaded file
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading("Research Dossier", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_paragraph(f"Research Run ID: {dossier.research_run_id}")
        doc.add_paragraph(f"Generated: {dossier.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        # Add executive summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(dossier.executive_summary)
        doc.add_page_break()

        # Add sections
        for section in dossier.get_sections_ordered():
            doc.add_heading(section.title, 1)
            doc.add_paragraph(section.content)
            doc.add_page_break()

        # Add citations if requested
        if include_citations:
            doc.add_heading("Citations", 1)
            doc.add_paragraph(dossier.citations_appendix)

        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Upload to MinIO
        filename = f"{base_filename}.docx"
        bucket = "dossiers"
        object_name = f"{dossier.research_run_id}/{filename}"

        await self.storage.upload_file(
            bucket_name=bucket,
            object_name=object_name,
            data=docx_buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return f"minio://{bucket}/{object_name}"

    async def _generate_pdf(
        self,
        dossier,
        base_filename: str,
        include_citations: bool,
        template_name: Optional[str],
    ) -> str:
        """
        Generate PDF file from dossier.

        Args:
            dossier: Dossier entity
            base_filename: Base filename without extension
            include_citations: Whether to include citations
            template_name: Optional template name

        Returns:
            MinIO path to uploaded file
        """
        # First generate DOCX in memory
        from docx import Document
        from docx2pdf import convert

        # Create temporary DOCX
        doc = Document()

        # Add title
        title = doc.add_heading("Research Dossier", 0)
        title.alignment = 1  # Center

        # Add metadata
        doc.add_paragraph(f"Research Run ID: {dossier.research_run_id}")
        doc.add_paragraph(f"Generated: {dossier.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        # Add executive summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(dossier.executive_summary)
        doc.add_page_break()

        # Add sections
        for section in dossier.get_sections_ordered():
            doc.add_heading(section.title, 1)
            doc.add_paragraph(section.content)
            doc.add_page_break()

        # Add citations if requested
        if include_citations:
            doc.add_heading("Citations", 1)
            doc.add_paragraph(dossier.citations_appendix)

        # Save DOCX to temp file
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            doc.save(tmp_docx.name)
            docx_path = tmp_docx.name

        # Convert to PDF
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
            pdf_path = tmp_pdf.name

        # Note: docx2pdf requires LibreOffice or Microsoft Word installed
        # For production, consider using reportlab or weasyprint instead
        try:
            convert(docx_path, pdf_path)
        except Exception as e:
            logger.warning(f"Failed to convert DOCX to PDF using docx2pdf: {e}")
            # Fallback: use alternate PDF generation (reportlab, etc.)
            logger.info("Using fallback PDF generation")
            pdf_buffer = await self._generate_pdf_fallback(dossier, include_citations)
            pdf_buffer.seek(0)

            # Upload fallback PDF
            filename = f"{base_filename}.pdf"
            bucket = "dossiers"
            object_name = f"{dossier.research_run_id}/{filename}"

            await self.storage.upload_file(
                bucket_name=bucket,
                object_name=object_name,
                data=pdf_buffer,
                content_type="application/pdf",
            )

            # Clean up temp files
            Path(docx_path).unlink(missing_ok=True)
            Path(pdf_path).unlink(missing_ok=True)

            return f"minio://{bucket}/{object_name}"

        # Read PDF
        with open(pdf_path, "rb") as f:
            pdf_data = BytesIO(f.read())

        # Upload to MinIO
        filename = f"{base_filename}.pdf"
        bucket = "dossiers"
        object_name = f"{dossier.research_run_id}/{filename}"

        await self.storage.upload_file(
            bucket_name=bucket,
            object_name=object_name,
            data=pdf_data,
            content_type="application/pdf",
        )

        # Clean up temp files
        Path(docx_path).unlink(missing_ok=True)
        Path(pdf_path).unlink(missing_ok=True)

        return f"minio://{bucket}/{object_name}"

    async def _generate_pdf_fallback(self, dossier, include_citations: bool) -> BytesIO:
        """
        Fallback PDF generation using reportlab.

        Args:
            dossier: Dossier entity
            include_citations: Whether to include citations

        Returns:
            BytesIO with PDF data
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            alignment=1,
        )
        story.append(Paragraph("Research Dossier", title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Metadata
        story.append(Paragraph(f"Research Run ID: {dossier.research_run_id}", styles["Normal"]))
        story.append(
            Paragraph(
                f"Generated: {dossier.generated_at.strftime('%Y-%m-%d %H:%M:%S')}",
                styles["Normal"],
            )
        )
        story.append(Spacer(1, 0.3 * inch))

        # Executive Summary
        story.append(Paragraph("Executive Summary", styles["Heading1"]))
        story.append(Paragraph(dossier.executive_summary, styles["Normal"]))
        story.append(PageBreak())

        # Sections
        for section in dossier.get_sections_ordered():
            story.append(Paragraph(section.title, styles["Heading1"]))
            story.append(Paragraph(section.content, styles["Normal"]))
            story.append(PageBreak())

        # Citations
        if include_citations:
            story.append(Paragraph("Citations", styles["Heading1"]))
            story.append(Paragraph(dossier.citations_appendix, styles["Normal"]))

        doc.build(story)
        buffer.seek(0)
        return buffer
