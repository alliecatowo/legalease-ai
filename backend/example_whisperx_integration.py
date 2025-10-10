"""
WhisperX Integration Example

This example demonstrates how to integrate the WhisperX pipeline
with the LegalEase backend's Celery task system.

This is a standalone example - the actual implementation is in
app/workers/tasks/transcription.py
"""

import os
import tempfile
from typing import Dict, Any
from io import BytesIO

from app.workers.pipelines import AudioPreprocessor, WhisperXPipeline
from app.core.database import SessionLocal
from app.core.minio_client import minio_client
from app.models.transcription import Transcription
from app.models.document import Document


def example_transcribe_document(document_id: int) -> Dict[str, Any]:
    """
    Example function showing how to transcribe a document.

    This demonstrates the complete workflow:
    1. Download audio from MinIO
    2. Preprocess to optimal format
    3. Transcribe with WhisperX
    4. Save results to database
    5. Upload transcripts to MinIO

    Args:
        document_id: ID of document to transcribe

    Returns:
        Dict with transcription results
    """
    db = SessionLocal()
    temp_dir = None

    try:
        # Get document from database
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise ValueError(f"Document {document_id} not found")

        print(f"Processing document: {document.filename}")

        # Create temporary directory
        temp_dir = tempfile.mkdtemp(prefix="transcription_")

        # Download audio from MinIO
        print("Downloading from MinIO...")
        audio_data = minio_client.download_file(document.file_path)

        original_path = os.path.join(temp_dir, "original" + os.path.splitext(document.filename)[1])
        with open(original_path, "wb") as f:
            f.write(audio_data)

        # Preprocess audio
        print("Preprocessing audio...")
        preprocessor = AudioPreprocessor(
            target_sample_rate=16000,
            target_channels=1,
        )

        preprocess_result = preprocessor.preprocess(
            input_file=original_path,
            output_file=os.path.join(temp_dir, "audio.wav")
        )

        wav_path = preprocess_result["output_path"]
        duration = preprocess_result["converted_metadata"]["duration"]

        # Transcribe with WhisperX
        print("Transcribing with WhisperX...")
        pipeline = WhisperXPipeline(
            model_name="large-v3",
            device="cuda",
            hf_token=os.getenv("HF_TOKEN")
        )

        result = pipeline.transcribe(
            audio_path=wav_path,
            enable_alignment=True,
            enable_diarization=True,  # Requires HF_TOKEN
            min_speakers=2,
            max_speakers=10,
        )

        print(f"Transcription complete: {len(result.segments)} segments")

        # Extract speaker information
        speakers_dict = {}
        if result.speakers:
            for speaker in result.speakers.get("unique_speakers", []):
                speakers_dict[speaker] = {
                    "id": speaker,
                    "segments_count": sum(
                        1 for seg in result.segments if seg.speaker == speaker
                    )
                }

        # Save to database
        print("Saving to database...")
        transcription = db.query(Transcription).filter(
            Transcription.document_id == document_id
        ).first()

        if not transcription:
            transcription = Transcription(
                document_id=document_id,
                format=os.path.splitext(document.filename)[1].lstrip('.'),
                duration=duration,
                speakers=list(speakers_dict.values()),
                segments=[seg.to_dict() for seg in result.segments]
            )
            db.add(transcription)
        else:
            transcription.duration = duration
            transcription.speakers = list(speakers_dict.values())
            transcription.segments = [seg.to_dict() for seg in result.segments]

        db.commit()

        # Export transcripts
        print("Exporting transcripts...")
        export_paths = export_transcripts(
            result=result,
            document_path=document.file_path,
            temp_dir=temp_dir,
        )

        print(f"Exported formats: {list(export_paths.keys())}")

        # Cleanup
        preprocessor.cleanup(wav_path)
        pipeline.cleanup()

        return {
            "status": "completed",
            "transcription_id": transcription.id,
            "duration": duration,
            "segments_count": len(result.segments),
            "speakers_count": result.num_speakers,
            "language": result.language,
            "export_paths": export_paths,
        }

    finally:
        # Cleanup temp directory
        if temp_dir and os.path.exists(temp_dir):
            import shutil
            shutil.rmtree(temp_dir)
        db.close()


def export_transcripts(
    result,
    document_path: str,
    temp_dir: str,
) -> Dict[str, str]:
    """
    Export transcription to multiple formats and upload to MinIO.

    Args:
        result: TranscriptionResult object
        document_path: Original document path in MinIO
        temp_dir: Temporary directory for exports

    Returns:
        Dict mapping format names to MinIO paths
    """
    import json
    from docx import Document
    from docx.shared import Pt, RGBColor

    export_paths = {}
    base_path = os.path.splitext(document_path)[0]

    # Export as JSON
    json_path = os.path.join(temp_dir, "transcription.json")
    with open(json_path, "w") as f:
        json.dump(result.to_dict(), f, indent=2)

    with open(json_path, "rb") as f:
        minio_path = f"{base_path}_transcription.json"
        minio_client.upload_file(
            file_data=BytesIO(f.read()),
            object_name=minio_path,
            content_type="application/json",
        )
        export_paths["json"] = minio_path

    # Export as plain text
    text_path = os.path.join(temp_dir, "transcription.txt")
    with open(text_path, "w") as f:
        for segment in result.segments:
            speaker = f"[{segment.speaker}] " if segment.speaker else ""
            timestamp = f"[{segment.start:.2f}s] "
            f.write(f"{speaker}{timestamp}{segment.text}\n")

    with open(text_path, "rb") as f:
        minio_path = f"{base_path}_transcription.txt"
        minio_client.upload_file(
            file_data=BytesIO(f.read()),
            object_name=minio_path,
            content_type="text/plain",
        )
        export_paths["txt"] = minio_path

    # Export as SRT subtitles
    srt_path = os.path.join(temp_dir, "transcription.srt")
    with open(srt_path, "w") as f:
        for idx, segment in enumerate(result.segments, 1):
            f.write(f"{idx}\n")
            f.write(f"{format_srt_time(segment.start)} --> {format_srt_time(segment.end)}\n")
            speaker = f"[{segment.speaker}] " if segment.speaker else ""
            f.write(f"{speaker}{segment.text}\n\n")

    with open(srt_path, "rb") as f:
        minio_path = f"{base_path}_transcription.srt"
        minio_client.upload_file(
            file_data=BytesIO(f.read()),
            object_name=minio_path,
            content_type="text/plain",
        )
        export_paths["srt"] = minio_path

    # Export as DOCX
    docx_path = os.path.join(temp_dir, "transcription.docx")
    doc = Document()

    # Title
    title = doc.add_heading("Transcription", 0)

    # Metadata
    doc.add_paragraph(f"Duration: {result.duration:.2f}s")
    doc.add_paragraph(f"Language: {result.language}")
    if result.num_speakers:
        doc.add_paragraph(f"Speakers: {result.num_speakers}")
    doc.add_paragraph("")

    # Segments
    current_speaker = None
    for segment in result.segments:
        # Add speaker heading if changed
        if segment.speaker and segment.speaker != current_speaker:
            current_speaker = segment.speaker
            speaker_para = doc.add_paragraph()
            speaker_run = speaker_para.add_run(f"\n{segment.speaker}")
            speaker_run.bold = True
            speaker_run.font.size = Pt(12)
            speaker_run.font.color.rgb = RGBColor(0, 0, 139)

        # Add text with timestamp
        para = doc.add_paragraph()
        time_run = para.add_run(f"[{segment.start:.2f}s] ")
        time_run.font.color.rgb = RGBColor(128, 128, 128)
        time_run.font.size = Pt(9)

        text_run = para.add_run(segment.text)
        text_run.font.size = Pt(11)

    doc.save(docx_path)

    with open(docx_path, "rb") as f:
        minio_path = f"{base_path}_transcription.docx"
        minio_client.upload_file(
            file_data=BytesIO(f.read()),
            object_name=minio_path,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        export_paths["docx"] = minio_path

    return export_paths


def format_srt_time(seconds: float) -> str:
    """Format seconds to SRT timestamp (HH:MM:SS,mmm)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def example_batch_transcription(document_ids: list) -> list:
    """
    Example of batch transcription processing.

    Args:
        document_ids: List of document IDs to transcribe

    Returns:
        List of transcription results
    """
    results = []

    # Initialize pipeline once for all documents
    pipeline = WhisperXPipeline(
        model_name="large-v3",
        device="cuda",
        hf_token=os.getenv("HF_TOKEN")
    )

    preprocessor = AudioPreprocessor()

    for doc_id in document_ids:
        try:
            print(f"\nProcessing document {doc_id}...")
            result = example_transcribe_document(doc_id)
            results.append(result)
        except Exception as e:
            print(f"Failed to transcribe document {doc_id}: {e}")
            results.append({
                "status": "failed",
                "document_id": doc_id,
                "error": str(e)
            })

    # Cleanup shared resources
    pipeline.cleanup()

    return results


def example_streaming_transcription():
    """
    Example of processing transcription with progress updates.

    This would be useful in a Celery task with self.update_state()
    """
    from app.workers.pipelines import WhisperXPipeline

    # Custom callback for progress updates
    def progress_callback(stage: str, progress: int):
        """Callback for progress updates."""
        print(f"[{progress}%] {stage}")
        # In Celery task, you would do:
        # self.update_state(state='PROCESSING', meta={'status': stage, 'progress': progress})

    # Simulate progress updates
    progress_callback("Downloading audio", 10)
    progress_callback("Preprocessing audio", 20)
    progress_callback("Loading Whisper model", 30)
    progress_callback("Transcribing...", 40)
    progress_callback("Aligning timestamps...", 60)
    progress_callback("Identifying speakers...", 80)
    progress_callback("Exporting results", 90)
    progress_callback("Complete", 100)


if __name__ == "__main__":
    """
    Example usage:

    # Single document
    result = example_transcribe_document(document_id=123)
    print(result)

    # Batch processing
    results = example_batch_transcription([123, 124, 125])
    for result in results:
        print(result)
    """

    print("WhisperX Integration Examples")
    print("=" * 60)
    print("\nThis file contains example code for integrating WhisperX")
    print("with the LegalEase backend.")
    print("\nKey functions:")
    print("  - example_transcribe_document(document_id)")
    print("  - example_batch_transcription(document_ids)")
    print("  - export_transcripts(result, document_path, temp_dir)")
    print("\nSee WHISPERX_SETUP.md for full documentation.")
